#!/usr/bin/env python3
"""
Terminal runner — Kombee AI Inventory Assistant
Minimal Tech theme: #ffffff / #f9fafb / #6366f1 accent
"""
import asyncio, json, sys, os, re
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from dotenv import load_dotenv
load_dotenv()
from backbone.orchestrator import AIOrchestrator

# ── palette (RGB tuples) ──────────────────────────────────────────────────────
WHITE      = (255, 255, 255)
OFF_WHITE  = (249, 250, 251)
GRAY_50    = (243, 244, 246)
GRAY_100   = (229, 231, 235)
GRAY_700   = ( 55,  65,  81)
GRAY_900   = ( 17,  24,  39)
ACCENT     = ( 99, 102, 241)   # #6366f1
ACCENT_SOFT= (224, 231, 255)   # #e0e7ff
RED        = (220,  38,  38)
AMBER      = (217, 119,   6)
GREEN      = ( 22, 163,  74)

# ── terminal input ────────────────────────────────────────────────────────────
def get_user_input() -> dict:
    print("\n" + "─" * 52)
    print("  Kombee  ·  Inventory Assistant")
    print("─" * 52)
    dish       = input("\n  Dish name        : ").strip()
    restaurant = input("  Restaurant name  : ").strip() or "My Restaurant"
    print("\n  Current stock  →  ingredient, quantity  (blank to finish)")
    stock = {}
    while True:
        line = input("  › ").strip()
        if not line:
            break
        parts = [p.strip() for p in line.split(",", 1)]
        if len(parts) == 2:
            stock[parts[0].lower()] = parts[1]
        else:
            print("  ⚠  format: ingredient, quantity")
    return {"dish": dish, "restaurant": restaurant, "stock": stock}

# ── parse JSON from LLM ───────────────────────────────────────────────────────
def parse_response(raw: str) -> dict:
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    if m:
        try:
            return json.loads(m.group())
        except json.JSONDecodeError:
            pass
    return {}

# ── terminal print ────────────────────────────────────────────────────────────
def print_order_list(dish: str, restaurant: str, data: dict):
    W = 52
    ingredients = data.get("ingredients", [])
    total_cost  = data.get("total_cost_inr", sum(i.get("cost_inr", 0) for i in ingredients))
    recipe      = data.get("recipe_steps", [])

    print("\n" + "═" * W)
    print(f"  {dish.upper()} — INGREDIENT ORDER LIST".center(W))
    print("═" * W)
    print(f"  Date        {datetime.now().strftime('%Y-%m-%d')}")
    print(f"  Restaurant  {restaurant}")

    groups = [
        ("🚨  CRITICAL — ORDER IMMEDIATELY", "CRITICAL"),
        ("⚠️   HIGH PRIORITY — ORDER SOON",  "HIGH"),
        ("📦  MEDIUM — ORDER THIS WEEK",     "MEDIUM"),
    ]
    counter = 1
    for heading, priority in groups:
        items = [i for i in ingredients if i.get("priority","").upper() == priority]
        if not items:
            continue
        print("\n" + "─" * W)
        print(f"  {heading}")
        print("─" * W)
        for item in items:
            print(f"\n  {counter}. {item['name']}")
            print(f"     Required       {item.get('required','N/A')}")
            print(f"     Current Stock  {item.get('current_stock','N/A')}")
            print(f"     Order Qty      {item.get('order_quantity','N/A')}")
            print(f"     Supplier       {item.get('supplier','N/A')}")
            print(f"     Cost           ₹{item.get('cost_inr','N/A')}")
            if item.get("note"):
                print(f"     Note           {item['note']}")
            counter += 1

    print("\n" + "═" * W)
    print(f"  TOTAL ORDER VALUE   ₹{total_cost}")
    print(f"  TOTAL ITEMS         {counter - 1}")
    print("═" * W)

    if recipe:
        print(f"\n  RECIPE · {dish.upper()}")
        print("─" * W)
        for step in recipe:
            print(f"  {step}")
        print("─" * W)

# ── PDF (Minimal Tech theme) ──────────────────────────────────────────────────
def save_pdf(dish: str, restaurant: str, data: dict):
    from fpdf import FPDF, XPos, YPos

    ingredients = data.get("ingredients", [])
    total_cost  = data.get("total_cost_inr", sum(i.get("cost_inr", 0) for i in ingredients))
    recipe      = data.get("recipe_steps", [])
    date_str    = datetime.now().strftime("%Y-%m-%d")

    pdf = FPDF()
    pdf.set_margins(14, 14, 14)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=16)
    W = pdf.w - pdf.l_margin - pdf.r_margin

    # ── header bar ──
    pdf.set_fill_color(*GRAY_900)
    pdf.rect(0, 0, pdf.w, 22, style="F")
    pdf.set_y(5)
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(*WHITE)
    pdf.cell(pdf.w, 10, f"{dish.upper()}  ·  INGREDIENT ORDER LIST", align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_y(24)

    # ── meta row ──
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(*GRAY_700)
    pdf.set_fill_color(*GRAY_50)
    pdf.rect(pdf.l_margin, pdf.get_y(), W, 8, style="F")
    pdf.cell(W/2, 8, f"  Date: {date_str}", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.cell(W/2, 8, f"Restaurant: {restaurant}  ", align="R",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(5)

    # ── section helper ──
    def section_header(label, color):
        pdf.set_fill_color(*color)
        pdf.set_text_color(*WHITE)
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(W, 7, f"  {label}", fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(*GRAY_900)
        pdf.ln(1)

    def kv(label, value, shade=False):
        pdf.set_fill_color(*(GRAY_50 if shade else WHITE))
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(38, 6, f"  {label}", fill=shade)
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(W - 38, 6, str(value), fill=shade,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    groups = [
        ("🚨  CRITICAL — ORDER IMMEDIATELY", "CRITICAL", RED),
        ("⚠️   HIGH PRIORITY — ORDER SOON",  "HIGH",     AMBER),
        ("📦  MEDIUM — ORDER THIS WEEK",     "MEDIUM",   ACCENT),
    ]
    counter = 1
    for heading, priority, color in groups:
        items = [i for i in ingredients if i.get("priority","").upper() == priority]
        if not items:
            continue
        section_header(heading, color)
        for item in items:
            # item title
            pdf.set_fill_color(*ACCENT_SOFT)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*ACCENT)
            pdf.cell(W, 7, f"  {counter}.  {item['name']}", fill=True,
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(*GRAY_900)
            shade = False
            for lbl, key in [("Required", "required"), ("Current Stock", "current_stock"),
                              ("Order Qty", "order_quantity"), ("Supplier", "supplier"),
                              ("Cost", None), ("Note", "note")]:
                val = f"Rs.{item.get('cost_inr','N/A')}" if key is None else item.get(key, "")
                if not val:
                    continue
                kv(lbl, val, shade)
                shade = not shade
            pdf.ln(2)
            counter += 1

    # ── totals bar ──
    pdf.set_fill_color(*GRAY_900)
    pdf.set_text_color(*WHITE)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(W/2, 8, f"  TOTAL ORDER VALUE:  Rs.{total_cost}", fill=True)
    pdf.cell(W/2, 8, f"TOTAL ITEMS:  {counter-1}  ", fill=True, align="R",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)

    # ── recipe ──
    if recipe:
        section_header(f"RECIPE  ·  {dish.upper()}", GREEN)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(*GRAY_700)
        for step in recipe:
            pdf.multi_cell(W, 5.5, f"  {step}")
            pdf.ln(0.5)

    fname = f"{dish.lower().replace(' ','_')}_order_list.pdf"
    pdf.output(fname)
    return fname


# ── Word (Minimal Tech theme) ─────────────────────────────────────────────────
def save_word(dish: str, restaurant: str, data: dict):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def hex_to_rgb(h):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.lstrip("#"))
        tcPr.append(shd)

    ingredients = data.get("ingredients", [])
    total_cost  = data.get("total_cost_inr", sum(i.get("cost_inr", 0) for i in ingredients))
    recipe      = data.get("recipe_steps", [])
    date_str    = datetime.now().strftime("%Y-%m-%d")

    doc = Document()

    # page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    # title
    t = doc.add_heading(f"{dish.upper()}  ·  INGREDIENT ORDER LIST", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = RGBColor(17, 24, 39)   # gray-900

    # meta
    meta = doc.add_paragraph(f"Date: {date_str}     Restaurant: {restaurant}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(55, 65, 81)

    section_colors = {
        "CRITICAL": "DC2626",
        "HIGH":     "D97706",
        "MEDIUM":   "6366F1",
    }
    section_labels = {
        "CRITICAL": "🚨  CRITICAL — ORDER IMMEDIATELY",
        "HIGH":     "⚠️   HIGH PRIORITY — ORDER SOON",
        "MEDIUM":   "📦  MEDIUM — ORDER THIS WEEK",
    }

    counter = 1
    for priority in ("CRITICAL", "HIGH", "MEDIUM"):
        items = [i for i in ingredients if i.get("priority","").upper() == priority]
        if not items:
            continue

        # section heading
        sh = doc.add_heading(section_labels[priority], level=2)
        sh.runs[0].font.color.rgb = RGBColor(*hex_to_rgb(section_colors[priority]))

        for item in items:
            # item name
            ih = doc.add_heading(f"{counter}.  {item['name']}", level=3)
            ih.runs[0].font.color.rgb = RGBColor(99, 102, 241)  # accent

            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Table Grid"
            tbl.columns[0].width = Cm(4.5)

            rows_data = [
                ("Required",      item.get("required", "N/A")),
                ("Current Stock", item.get("current_stock", "N/A")),
                ("Order Qty",     item.get("order_quantity", "N/A")),
                ("Supplier",      item.get("supplier", "N/A")),
                ("Cost",          f"Rs.{item.get('cost_inr','N/A')}"),
            ]
            if item.get("note"):
                rows_data.append(("Note", item["note"]))

            for idx, (lbl, val) in enumerate(rows_data):
                r = tbl.add_row().cells
                r[0].text = lbl
                r[1].text = str(val)
                r[0].paragraphs[0].runs[0].bold = True
                r[0].paragraphs[0].runs[0].font.size = Pt(8.5)
                r[1].paragraphs[0].runs[0].font.size = Pt(8.5)
                bg = "F3F4F6" if idx % 2 == 0 else "FFFFFF"
                set_cell_bg(r[0], bg)
                set_cell_bg(r[1], bg)

            doc.add_paragraph()
            counter += 1

    # totals
    tot = doc.add_heading(
        f"TOTAL ORDER VALUE: Rs.{total_cost}     TOTAL ITEMS: {counter-1}", level=2)
    tot.runs[0].font.color.rgb = RGBColor(17, 24, 39)

    # recipe
    if recipe:
        rh = doc.add_heading(f"RECIPE  ·  {dish.upper()}", level=1)
        rh.runs[0].font.color.rgb = RGBColor(22, 163, 74)
        for step in recipe:
            p = doc.add_paragraph(step, style="List Number")
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGBColor(55, 65, 81)

    fname = f"{dish.lower().replace(' ','_')}_order_list.docx"
    doc.save(fname)
    return fname


# ── main ──────────────────────────────────────────────────────────────────────
async def main():
    inp = get_user_input()
    dish, restaurant, stock = inp["dish"], inp["restaurant"], inp["stock"]

    print("\n  Processing…\n")
    orchestrator = AIOrchestrator()
    result = await orchestrator.process(
        query=f"Generate ingredient order list for {dish}",
        context={
            "dish": dish,
            "restaurant_name": restaurant,
            "current_stock": stock,
            "user_id": "terminal_user",
            "session_id": "session_001",
        },
    )

    data = parse_response(result["response"])
    if not data:
        print("  Could not parse structured response. Raw output:\n")
        print(result["response"])
        return

    print_order_list(dish, restaurant, data)

    pdf_file  = save_pdf(dish, restaurant, data)
    word_file = save_word(dish, restaurant, data)

    print(f"\n  ✓  PDF   →  {pdf_file}")
    print(f"  ✓  Word  →  {word_file}\n")


if __name__ == "__main__":
    if sys.platform == "win32":
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    asyncio.run(main())
